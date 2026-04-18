"""DOCX protocol parser implementation using python-docx.

This module provides a production-grade DOCX parser that extracts structured
protocol data from Microsoft Word documents using the python-docx library.

Features:
- Primary: LLM-powered extraction for robust handling of any format
- Fallback: Regex-based extraction for when LLM is unavailable
"""

from __future__ import annotations

import re
from io import BytesIO
from typing import Any, Dict, List, Optional

import docx

from app.llm.base_llm import BaseLLM
from app.parsing.base_parser import BaseParser
from app.parsing.llm_extraction_service import LLMExtractionService
from app.core.exceptions import FileParsingError
from app.core.logging import get_logger

logger = get_logger(__name__)


class DOCXParser(BaseParser):
    """DOCX protocol parser using python-docx for text extraction.
    
    This parser extracts text from Microsoft Word documents and uses LLM-powered
    extraction as the primary method for robust handling of any format. Falls back
    to regex-based extraction if LLM is unavailable.
    
    Features:
        - Paragraph and table text extraction with python-docx
        - LLM-powered structured data extraction (primary)
        - Regex-based extraction (fallback)
        - Handles any DOCX format or layout
        - Robust error handling and logging
    """
    
    def __init__(self, llm: Optional[BaseLLM] = None) -> None:
        """Initialize the DOCX parser.
        
        Args:
            llm: Optional LLM provider for intelligent extraction.
                 If None, falls back to regex-based extraction.
        """
        super().__init__()
        self._llm = llm
        self._extraction_service = LLMExtractionService(llm) if llm else None
    
    @property
    def supported_extensions(self) -> List[str]:
        """Return supported file extensions."""
        return ['.docx', '.DOCX', '.doc', '.DOC']
    
    @property
    def supported_mime_types(self) -> List[str]:
        """Return supported MIME types."""
        return [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]
    
    @property
    def parser_name(self) -> str:
        """Return parser name."""
        return "DOCXParser"
    
    def parse(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Parse DOCX file and extract structured protocol data.
        
        Args:
            file_bytes: The DOCX file content as bytes.
            filename: The original filename.
        
        Returns:
            Dictionary containing structured protocol data.
        
        Raises:
            FileParsingError: When the DOCX cannot be parsed.
        """
        self.validate_input(file_bytes, filename)
        
        self._logger.info(
            "[%s] Parsing DOCX — filename=%s, size=%d bytes",
            self.parser_name,
            filename,
            len(file_bytes)
        )
        
        try:
            # Extract text from DOCX
            text = self._extract_text_from_docx(file_bytes)
            
            if not text or len(text.strip()) < 10:
                raise FileParsingError(
                    "DOCX appears to be empty or contains no extractable text.",
                    details={"filename": filename, "text_length": len(text)}
                )
            
            # Extract structured data using LLM (primary) or regex (fallback)
            if self._extraction_service:
                self._logger.info(
                    "[%s] Using LLM-powered extraction for robust parsing",
                    self.parser_name
                )
                protocol_data = self._extraction_service.extract_protocol_data(text, filename)
            else:
                self._logger.info(
                    "[%s] Using regex-based extraction (LLM not available)",
                    self.parser_name
                )
                protocol_data = self._extract_protocol_from_text(text, filename)
            
            # Add metadata
            protocol_data.update({
                "source": "docx",
                "filename": filename,
                "raw_text": text,
                "raw_text_length": len(text)
            })
            
            self._logger.info(
                "[%s] Successfully parsed DOCX — extracted %d characters, protocol_id=%s",
                self.parser_name,
                len(text),
                protocol_data.get("protocol_id", "UNKNOWN")
            )
            
            return protocol_data
            
        except FileParsingError:
            raise
        except Exception as e:
            self._logger.error(
                "[%s] Failed to parse DOCX: %s",
                self.parser_name,
                str(e),
                exc_info=True
            )
            raise FileParsingError(
                f"Failed to parse DOCX: {e}",
                details={"filename": filename, "parser": self.parser_name}
            ) from e
    
    def _extract_text_from_docx(self, file_bytes: bytes) -> str:
        """Extract all text from DOCX file bytes.
        
        Extracts text from both paragraphs and tables to ensure comprehensive
        content extraction.
        
        Args:
            file_bytes: The DOCX content as bytes.
        
        Returns:
            Extracted text from document.
        
        Raises:
            Exception: If DOCX cannot be opened or read.
        """
        text = ""
        
        try:
            # Create file-like object from bytes
            docx_stream = BytesIO(file_bytes)
            
            # Open the document
            doc = docx.Document(docx_stream)
            
            # Extract text from paragraphs
            paragraph_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables
            table_count = 0
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text += row_text + "\n"
                table_count += 1
            
            self._logger.debug(
                "[%s] Extracted text from %d paragraphs and %d tables",
                self.parser_name,
                paragraph_count,
                table_count
            )
            
        except Exception as e:
            self._logger.error("[%s] Error reading DOCX: %s", self.parser_name, str(e))
            raise
        
        return text
    
    def _extract_protocol_from_text(self, text: str, filename: str) -> Dict[str, Any]:
        """Extract structured protocol data from text using heuristics.
        
        This method uses regex patterns to identify key protocol fields.
        For production use with complex protocols, consider:
        1. Using an LLM for more accurate extraction
        2. Training a custom NER model
        3. Using template-based extraction for known formats
        
        Args:
            text: The extracted text from DOCX.
            filename: The original filename.
        
        Returns:
            Dictionary with structured protocol data.
        """
        # Start with default structure
        protocol_data = self._create_default_protocol_data(filename, "docx")
        
        lines = text.split('\n')
        
        # Extract protocol ID
        protocol_id = self._extract_protocol_id(lines)
        if protocol_id:
            protocol_data["protocol_id"] = protocol_id
        
        # Extract phase
        phase = self._extract_phase(lines)
        if phase:
            protocol_data["phase"] = phase
        
        # Extract therapeutic area
        therapeutic_area = self._extract_therapeutic_area(lines)
        if therapeutic_area:
            protocol_data["therapeutic_area"] = therapeutic_area
        
        # Extract target enrollment
        enrollment = self._extract_target_enrollment(lines)
        if enrollment:
            protocol_data["target_enrollment"] = enrollment
        
        # Extract title
        title = self._extract_title(lines)
        if title:
            protocol_data["title"] = title
        
        # Extract criteria sections
        protocol_data["inclusion_criteria"] = self._extract_section(
            text,
            ["inclusion criteria", "key inclusion criteria", "inclusion"]
        )
        
        protocol_data["exclusion_criteria"] = self._extract_section(
            text,
            ["exclusion criteria", "key exclusion criteria", "exclusion"]
        )
        
        # Extract endpoints
        protocol_data["primary_endpoints"] = self._extract_section(
            text,
            ["primary endpoint", "primary endpoints", "primary outcome"]
        )
        
        return protocol_data
    
    def _extract_protocol_id(self, lines: List[str]) -> str:
        """Extract protocol ID from text lines."""
        for line in lines:
            line = line.strip()
            match = re.search(
                r'protocol\s*(?:id|number|no\.?)[:\s]+([A-Z0-9-]{3,})',
                line,
                re.IGNORECASE
            )
            if match:
                return match.group(1)
        return ""
    
    def _extract_phase(self, lines: List[str]) -> str:
        """Extract study phase from text lines."""
        for line in lines:
            line = line.strip()
            match = re.search(r'phase\s*(I{1,3}|IV|1|2|3|4)', line, re.IGNORECASE)
            if match:
                phase = match.group(1).upper()
                phase_map = {'1': 'I', '2': 'II', '3': 'III', '4': 'IV'}
                return f"Phase {phase_map.get(phase, phase)}"
        return ""
    
    def _extract_therapeutic_area(self, lines: List[str]) -> str:
        """Extract therapeutic area from text lines."""
        for line in lines:
            line = line.strip()
            match = re.search(
                r'(?:therapeutic\s*area|indication|disease\s*area)[:\s]+([A-Za-z\s]+)',
                line,
                re.IGNORECASE
            )
            if match:
                area = match.group(1).strip()
                words = area.split()[:3]
                return ' '.join(words)
        return ""
    
    def _extract_target_enrollment(self, lines: List[str]) -> int:
        """Extract target enrollment number from text lines."""
        for line in lines:
            line = line.strip()
            match = re.search(
                r'(?:target\s*enrollment|sample\s*size|number\s*of\s*patients)[:\s]+(\d+)',
                line,
                re.IGNORECASE
            )
            if match:
                return int(match.group(1))
        return 0
    
    def _extract_title(self, lines: List[str]) -> str:
        """Extract protocol title from text lines."""
        for line in lines[:20]:
            line = line.strip()
            if len(line) < 20:
                continue
            if line.isupper() or line.istitle():
                if any(keyword in line.lower() for keyword in ['page', 'confidential', 'protocol']):
                    continue
                return line[:200]
        return ""
    
    def _extract_section(self, text: str, section_headers: List[str]) -> List[str]:
        """Extract bullet points or numbered items from a section.
        
        Args:
            text: The full text to search.
            section_headers: List of possible section header names.
        
        Returns:
            List of extracted items (up to 10).
        """
        section_text = ""
        for header in section_headers:
            pattern = rf"{header}[:\s]*\n([\s\S]*?)(?=\n\s*[A-Z][A-Za-z\s]{{10,}}:|\n\s*\d+\.\s*[A-Z]|$)"
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                section_text = match.group(1)
                break
        
        if not section_text:
            return []
        
        items = []
        
        # Extract bullet points
        bullets = re.findall(r'^\s*[•·\-\*]\s*(.+)$', section_text, re.MULTILINE)
        items.extend(bullets)
        
        # Extract numbered items
        numbered = re.findall(r'^\s*\d+\.\s*(.+)$', section_text, re.MULTILINE)
        items.extend(numbered)
        
        # If no structured items found, split by newlines
        if not items:
            lines = [line.strip() for line in section_text.split('\n') if line.strip()]
            items = [line for line in lines if len(line) > 15][:10]
        
        # Clean up items
        cleaned_items = []
        for item in items[:10]:
            item = item.strip()
            item = re.sub(r'^[•·\-\*\d\.]+\s*', '', item)
            item = item.rstrip('.,;')
            if item and len(item) > 10:
                cleaned_items.append(item)
        
        return cleaned_items
